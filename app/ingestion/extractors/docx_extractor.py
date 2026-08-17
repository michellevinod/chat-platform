from pathlib import Path

from docx import Document

from app.ingestion.extractors.base_extractor import BaseExtractor
from app.ingestion.extractors.raw_models import (
    RawBlock,
    RawDocument,
    RawPage,
    RawTableBlock,
    RawTextBlock,
)
from app.models.enums import BlockType


class DOCXExtractor(BaseExtractor):
    """
    Extracts text and tables from Microsoft Word (.docx) documents.
    """

    def extract(
        self,
        file_path: Path,
    ) -> RawDocument:

        document = Document(file_path)

        blocks: list[RawBlock] = []

        block_number = 0

        for paragraph in document.paragraphs:

            text = paragraph.text.strip()

            if not text:
                continue

            blocks.append(
                RawTextBlock(
                    page_number=1,
                    block_number=block_number,
                    block_type=BlockType.TEXT,
                    bbox=(0.0, 0.0, 0.0, 0.0),
                    text=text,
                )
            )

            block_number += 1

        # Extract tables
        for tab_idx, table in enumerate(document.tables):
            rows_data = []
            for row in table.rows:
                row_vals = [cell.text.strip() for cell in row.cells]
                rows_data.append(row_vals)

            if rows_data and len(rows_data) > 0:
                headers = rows_data[0]
                lines = []
                lines.append("| " + " | ".join(headers) + " |")
                lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
                for r in rows_data[1:]:
                    lines.append("| " + " | ".join(r) + " |")
                md_table = "\n".join(lines)

                blocks.append(
                    RawTableBlock(
                        page_number=1,
                        block_number=block_number,
                        block_type=BlockType.TABLE,
                        bbox=(0.0, 0.0, 0.0, 0.0),
                        markdown=md_table,
                        rows=rows_data,
                        headers=headers,
                        caption=f"Table {tab_idx + 1}",
                    )
                )
                block_number += 1

        page = RawPage(
            page_number=1,
            blocks=blocks,
        )

        return RawDocument(
            file_name=file_path.name,
            total_pages=1,
            pages=[page],
        )